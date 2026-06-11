import os
import re
import json
import smtplib
import requests
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date
import anthropic
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load secrets from Streamlit Cloud or fall back to environment variables
def get_secret(key):
    try:
        import streamlit as st
        return st.secrets[key]
    except:
        return os.getenv(key)

# Initialize clients
claude_client = anthropic.Anthropic(api_key=get_secret("ANTHROPIC_API_KEY"))
grok_client = OpenAI(api_key=get_secret("XAI_API_KEY"), base_url="https://api.x.ai/v1")
openai_client = OpenAI(api_key=get_secret("OPENAI_API_KEY"))
github_client = OpenAI(api_key=get_secret("GITHUB_TOKEN"), base_url="https://models.github.ai/inference")
perplexity_client = OpenAI(api_key=get_secret("PERPLEXITY_API_KEY"), base_url="https://api.perplexity.ai")
SENDER_EMAIL = "kirkkeller@gmail.com"
GMAIL_APP_PASSWORD = get_secret("GMAIL_APP_PASSWORD")

# Initialize xAI native client for web search
from xai_sdk import Client as XAIClient
xai_client = XAIClient(api_key=get_secret("XAI_API_KEY"))

# Routing prompt
ROUTING_PROMPT = """You are an AI task router. Analyze the task and respond with ONLY one of these words:
- claude: reasoning, analysis, writing, summarization, formatting, general questions
- grok: current events, live web search, news, recent information
- openai: creative writing, broad general knowledge tasks
- github: ANY coding task, code generation, code review, debugging, technical documentation, programming questions, software development, writing functions or scripts
- research: search for contact information, company details, recruiter info, business research, or any research requiring cited sources then email results
- multi: ONLY for searching job boards for employment opportunities, job openings, or job postings for a candidate

Respond with ONLY one word. Nothing else."""


def route_task(user_input):
    """Ask Claude to decide which AI should handle this task."""
    response = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=10,
        messages=[
            {"role": "user", "content": f"{ROUTING_PROMPT}\n\nTask: {user_input}"}
        ]
    )
    return response.content[0].text.strip().lower()


def build_verified_search_urls(job_title, company):
    """Build guaranteed working search URLs for a job at a company."""
    title_encoded = requests.utils.quote(job_title)
    company_encoded = requests.utils.quote(company)
    query = f"{title_encoded}+{company_encoded}"
    return {
        "LinkedIn": f"https://www.linkedin.com/jobs/search/?keywords={query}&f_WT=2",
        "Indeed": f"https://www.indeed.com/jobs?q={query}&l=Remote",
        "Glassdoor": f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={query}"
    }


def validate_url(url):
    """Check if a URL returns a valid response."""
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        response = requests.head(url, headers=headers, timeout=5, allow_redirects=True)
        return response.status_code in [200, 301, 302]
    except:
        return False


def parse_jobs_from_grok(raw_text):
    """Ask Claude to parse Grok's response into structured job data."""
    response = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        messages=[
            {"role": "user", "content": "Parse the following job search results into a JSON array. Each job should have these exact fields: title, company, salary, location, description. Return ONLY a valid JSON array with no other text, no markdown, no backticks. If you cannot parse any jobs, return an empty array: []\n\nRaw results:\n" + raw_text}
        ]
    )
    raw_response = response.content[0].text.strip()
    print(f"\n--- DEBUG: Claude parse response (first 300 chars) ---")
    print(raw_response[:300])
    print("--- END DEBUG ---\n")
    try:
        clean = raw_response.replace("```json", "").replace("```", "").strip()
        jobs = json.loads(clean)
        return jobs if isinstance(jobs, list) else []
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        return []


def search_with_grok(user_input, exclude_companies=None):
    """Use Grok with live web search tool enabled."""
    from xai_sdk.chat import user as xai_user
    from xai_sdk.tools import web_search as grok_web_search

    exclude_text = ""
    if exclude_companies:
        exclude_text = f"\nDo NOT include positions from: {', '.join(exclude_companies)}"

    search_query = (
        "Search the web right now for real current remote job postings. "
        "Find at least 10 remote Senior Sourcing or Procurement Manager positions paying $140,000-$175,000+. "
        "Search LinkedIn, Indeed, Glassdoor for these titles: "
        "Senior Sourcing Manager, Director of Sourcing, Strategic Sourcing Manager, "
        "Senior Procurement Manager, Global Sourcing Manager, Contract Manager, Vendor Manager. "
        "Requirements: Remote only, US-based, salary $140k-$175k+. "
        + exclude_text +
        "\nFor each job return: Job Title, Company, Salary, Location, Description (2-3 sentences)."
    )

    chat = xai_client.chat.create(model="grok-3", tools=[grok_web_search()])
    chat.append(xai_user(search_query))
    response = chat.sample()

    if hasattr(response, 'message'):
        result = response.message.content
    elif hasattr(response, 'choices'):
        result = response.choices[0].message.content
    elif hasattr(response, 'content'):
        result = response.content
    else:
        result = str(response)

    print(f"\n--- DEBUG: Grok raw response (first 500 chars) ---")
    print(result[:500])
    print("--- END DEBUG ---\n")
    return result


def research_with_grok(user_input):
    """Use Grok to research contact information and business details."""
    from xai_sdk.chat import user as xai_user
    from xai_sdk.tools import web_search as grok_web_search

    chat = xai_client.chat.create(model="grok-3", tools=[grok_web_search()])
    chat.append(xai_user(user_input))
    response = chat.sample()

    if hasattr(response, 'message'):
        result = response.message.content
    elif hasattr(response, 'choices'):
        result = response.choices[0].message.content
    elif hasattr(response, 'content'):
        result = response.content
    else:
        result = str(response)

    print(f"\n--- DEBUG: Grok research response (first 500 chars) ---")
    print(result[:500])
    print("--- END DEBUG ---\n")
    return result


def research_with_perplexity(user_input):
    """Use Perplexity for research tasks with cited sources."""
    response = perplexity_client.chat.completions.create(
        model="sonar",
        messages=[
            {"role": "system", "content": "You are a research assistant. Search the web thoroughly and provide detailed, accurate results with citations and sources for every claim."},
            {"role": "user", "content": user_input}
        ]
    )
    result = response.choices[0].message.content
    print(f"\n--- DEBUG: Perplexity response (first 500 chars) ---")
    print(result[:500])
    print("--- END DEBUG ---\n")
    return result


def research_with_both(user_input):
    """Use Grok for discovery, Perplexity for verification."""
    print("  Discovery search with Grok...")
    grok_results = research_with_grok(user_input)
    print("  Verification search with Perplexity...")
    perplexity_results = research_with_perplexity(user_input)
    return f"GROK FINDINGS:\n{grok_results}\n\nPERPLEXITY VERIFIED:\n{perplexity_results}"


def create_word_doc(content, title="Research Results"):
    """Convert research text content into a formatted Word document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Process content line by line
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x9E)

        elif line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                run.font.size = Pt(11)
                if j % 2 == 1:
                    run.font.bold = True

        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            run = p.add_run(text)
            run.font.size = Pt(11)

        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', line)
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    if j % 2 == 1:
                        run.font.bold = True

    filepath = os.path.join(tempfile.gettempdir(), f"research_{title[:30].replace(' ', '_').replace('/', '_')}.docx")
    doc.save(filepath)
    return filepath


def format_with_claude(jobs_with_links, user_input):
    """Use Claude to format validated jobs into a clean email."""
    jobs_text = json.dumps(jobs_with_links, indent=2)
    prompt = (
        "Format these verified job search results into a clean email. "
        "Each job has been verified and includes working search links. "
        "Format each job clearly with all details. "
        "Note that links go to job board searches for that specific role at that company. "
        "End with a brief note about search criteria used.\n\n"
        f"Original request summary: {user_input[:200]}\n\nJobs data:\n{jobs_text}"
    )
    response = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text


def format_research_with_claude(raw_results, user_input):
    """Use Claude to format research results as a proper document."""
    prompt = (
        "Format the following research findings into a clear professional document. "
        "This is NOT a job search - it is a research task.\n\n"
        "Format it as a proper briefing document with:\n"
        "- Clear sections and headers using markdown (## for sections, ### for subsections)\n"
        "- Bullet points where appropriate\n"
        "- Bold key terms using **term** markdown\n"
        "- Professional tone appropriate for the audience described in the request\n"
        "- Sources cited where available\n\n"
        f"Original request: {user_input}\n\n"
        f"Research findings:\n{raw_results}"
    )
    response = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text


def send_email(to_address, subject, body):
    """Send plain text email via Gmail SMTP."""
    msg = MIMEMultipart()
    msg['From'] = SENDER_EMAIL
    msg['To'] = to_address
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
        server.login(SENDER_EMAIL, GMAIL_APP_PASSWORD)
        server.send_message(msg)


def send_email_with_attachment(to_address, subject, body, attachment_path):
    """Send email with a Word doc attachment via Gmail SMTP."""
    msg = MIMEMultipart()
    msg['From'] = SENDER_EMAIL
    msg['To'] = to_address
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    with open(attachment_path, 'rb') as f:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        filename = os.path.basename(attachment_path)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
        server.login(SENDER_EMAIL, GMAIL_APP_PASSWORD)
        server.send_message(msg)


def handle_research_task(user_input):
    """Handle research tasks - search web, create Word doc, email as attachment."""
    print("Searching with Grok + Perplexity...")
    raw_results = research_with_both(user_input)

    print("Formatting results with Claude...")
    formatted_results = format_research_with_claude(raw_results, user_input)

    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', user_input)
    to_address = email_match.group(0) if email_match else "kirkkeller@gmail.com"

    subject_match = re.search(r'subject[:\s]+"?([^"\n]+)"?', user_input, re.IGNORECASE)
    subject = subject_match.group(1).strip() if subject_match else f"Research Results - {date.today().strftime('%B %d, %Y')}"

    print("Creating Word document...")
    doc_path = create_word_doc(formatted_results, subject)

    print(f"Sending email with attachment to {to_address}...")
    email_body = f"Please find the research results attached as a Word document.\n\nDocument: {os.path.basename(doc_path)}\nDate: {date.today().strftime('%B %d, %Y')}"
    send_email_with_attachment(
        to_address=to_address,
        subject=subject,
        body=email_body,
        attachment_path=doc_path
    )

    # Clean up temp file
    try:
        os.remove(doc_path)
    except:
        pass

    return f"Done! Research results emailed as Word document to {to_address}\n\nPreview:\n{formatted_results[:500]}..."


def handle_multi_task(user_input, target_valid=10, max_attempts=5):
    """Handle tasks requiring search + format + email with recursive validation."""
    exclude_companies = []
    if "First Citizens" in user_input:
        exclude_companies = ["First Citizens", "First Citizens Bank & Trust"]

    validated_jobs = []
    attempt = 0
    seen_companies = set()

    while len(validated_jobs) < target_valid and attempt < max_attempts:
        attempt += 1
        print(f"Step {attempt}: Searching for jobs (attempt {attempt}/{max_attempts})...")

        raw_results = search_with_grok(user_input, exclude_companies)

        print(f"   Parsing job listings...")
        jobs = parse_jobs_from_grok(raw_results)

        print(f"   Found {len(jobs)} listings, building and validating links...")

        for job in jobs:
            company_key = job.get("company", "").lower().strip()
            if company_key in seen_companies:
                continue
            if any(exc.lower() in company_key for exc in exclude_companies):
                continue

            search_urls = build_verified_search_urls(job.get("title", ""), job.get("company", ""))

            verified_urls = {}
            for platform, url in search_urls.items():
                if validate_url(url):
                    verified_urls[platform] = url

            if verified_urls:
                job["search_links"] = verified_urls
                job["link_status"] = "Verified"
                validated_jobs.append(job)
                seen_companies.add(company_key)
                print(f"   OK: {job.get('company')} - {job.get('title')}")
            else:
                print(f"   Skipping {job.get('company')} - no valid links found")

            if len(validated_jobs) >= target_valid:
                break

        print(f"   Have {len(validated_jobs)}/{target_valid} verified jobs so far...")

    if not validated_jobs:
        return "No verified job listings found after maximum search attempts."

    print(f"\nFormatting {len(validated_jobs)} verified results with Claude...")
    formatted_results = format_with_claude(validated_jobs, user_input)

    print("Sending email...")
    send_email(
        to_address="cindyrkeller@gmail.com",
        subject=f"Job Search Results - {date.today().strftime('%B %d, %Y')}",
        body=formatted_results
    )

    return f"Done! {len(validated_jobs)} verified jobs emailed to cindyrkeller@gmail.com\n\nPreview:\n{formatted_results[:500]}..."


def run_task(ai, user_input):
    """Send the task to the chosen AI and return the result."""
    if ai == "research":
        return handle_research_task(user_input)

    if ai == "multi":
        return handle_multi_task(user_input)

    elif ai == "claude":
        response = claude_client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1024,
            messages=[{"role": "user", "content": user_input}]
        )
        return response.content[0].text

    elif ai == "grok":
        response = grok_client.chat.completions.create(
            model="grok-3",
            messages=[
                {"role": "system", "content": "You are a research assistant with live web search. Search the web and return detailed results with links and sources."},
                {"role": "user", "content": user_input}
            ]
        )
        return response.choices[0].message.content

    elif ai == "openai":
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": user_input}]
        )
        return response.choices[0].message.content

    elif ai == "github":
        response = github_client.chat.completions.create(
            model="openai/gpt-4.1",
            messages=[
                {"role": "system", "content": "You are an expert software engineer and technical assistant. Help with code generation, code review, debugging, and technical documentation."},
                {"role": "user", "content": user_input}
            ]
        )
        return response.choices[0].message.content


def main():
    print("AI Router ready. Type 'quit' or paste your prompt, then press Enter twice when done.\n")
    while True:
        print("Your task (press Enter twice when done):")
        lines = []
        while True:
            line = input()
            if line.strip() == "###END###":
                break
            else:
                lines.append(line)
        user_input = "\n".join(lines).strip()

        if user_input.lower() == "quit":
            break
        if not user_input:
            continue

        print("\nRouting your task...")
        chosen_ai = route_task(user_input)
        print(f"Routing to: {chosen_ai.upper()}\n")

        result = run_task(chosen_ai, user_input)
        print(f"Result:\n{result}\n")
        print("-" * 50 + "\n")


if __name__ == "__main__":
    main()