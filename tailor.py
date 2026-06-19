"""
tailor.py — Job tailoring engine for Cindy's job search app.

Given a job posting (URL or pasted text), this module:
  1. Fetches the job description (URL first, paste fallback)
  2. Reads Cindy's master resume and cover letter from ~/cindy/resume/
  3. Calls Claude to tailor both documents
  4. Saves tailored .docx files to ~/cindy/output/<Company>_<Date>/

Usage:
    from tailor import tailor_for_job
    result = tailor_for_job(job_id, url, pasted_text)
"""

import os
import re
import requests
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt
import anthropic

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

HOME           = Path.home()
RESUME_DIR     = HOME / "cindy" / "resume"
OUTPUT_DIR     = HOME / "cindy" / "output"
MASTER_RESUME  = RESUME_DIR / "Cindy_Keller_Resume_v2.docx"
MASTER_COVER   = RESUME_DIR / "Cindy_Keller_Cover_Letter_Master.docx"

# ---------------------------------------------------------------------------
# Anthropic client
# ---------------------------------------------------------------------------

def _get_client() -> anthropic.Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not set in environment")
    return anthropic.Anthropic(api_key=api_key)


# ---------------------------------------------------------------------------
# Job description fetching
# ---------------------------------------------------------------------------

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}

def fetch_job_description(url: str) -> tuple[str, str]:
    """
    Attempt to fetch and extract job description text from a URL.

    Returns:
        (text, status) where status is 'ok', 'blocked', or 'error'
    """
    if not url or not url.startswith("http"):
        return "", "no_url"

    try:
        response = requests.get(url, headers=HEADERS, timeout=15)
        if response.status_code == 200:
            # Strip HTML tags — simple but effective for job boards
            text = re.sub(r'<[^>]+>', ' ', response.text)
            text = re.sub(r'\s+', ' ', text).strip()
            # Keep only the most relevant portion (cap at 8000 chars)
            if len(text) > 8000:
                text = text[:8000]
            if len(text) < 200:
                return "", "blocked"
            return text, "ok"
        elif response.status_code in (403, 429):
            return "", "blocked"
        else:
            return "", f"error_{response.status_code}"
    except Exception as e:
        return "", f"error: {str(e)[:100]}"


# ---------------------------------------------------------------------------
# Document reading
# ---------------------------------------------------------------------------

def read_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file."""
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Claude tailoring
# ---------------------------------------------------------------------------

RESUME_SYSTEM = """You are an expert resume writer specializing in Supply Chain, 
Sourcing, and Procurement roles. You tailor resumes to match job postings precisely, 
mirroring the employer's language and keywords while keeping all content truthful.

Rules:
- Mirror the job posting's exact terminology and keywords
- Reorder bullet points to lead with most relevant experience
- Adjust the professional summary to speak directly to this role
- Keep all facts, dates, companies, and metrics exactly as provided
- Do not invent experience or qualifications
- Output the complete tailored resume as plain text, preserving section structure
- Use the same section headers as the original resume"""

COVER_SYSTEM = """You are an expert cover letter writer specializing in Supply Chain, 
Sourcing, and Procurement roles. You write concise, targeted cover letters that 
speak directly to the employer's needs.

Rules:
- Output the COMPLETE document in this exact structure, one item per line:
    Line 1:  Cindy Keller
    Line 2:  cindyrkeller@gmail.com  |  (919) 349-4900  |  Piney Creek, NC  |  Remote
    Line 3:  (blank)
    Line 4:  Today's date (e.g. June 19, 2026)
    Line 5:  Dear Hiring Team,
    Line 6:  Opening paragraph — hook and fit, 2-3 sentences
    Line 7:  Middle paragraph — proof and experience, 2-3 sentences
    Line 8:  Closing paragraph — close and call to action, 2 sentences
    Line 9:  Thank you for your time and consideration.
    Line 10: Sincerely,
    Line 11: Cindy Keller
- Mirror the job posting's language and priorities
- Reference the company name and role title specifically
- Keep body paragraphs under 100 words each
- Professional but warm tone
- Output ONLY these 11 lines, nothing else — no customization notes"""


def tailor_resume_with_claude(
    master_resume: str,
    job_description: str,
    company: str,
    title: str
) -> str:
    """Call Claude to tailor the resume for this job posting."""
    client = _get_client()

    prompt = f"""Please tailor this resume for the following job posting.

JOB POSTING:
Company: {company}
Title: {title}

{job_description}

---

MASTER RESUME:
{master_resume}

---

Output the complete tailored resume as plain text."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=RESUME_SYSTEM,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def tailor_cover_letter_with_claude(
    master_cover: str,
    job_description: str,
    company: str,
    title: str
) -> str:
    """Call Claude to write a targeted cover letter for this job posting."""
    client = _get_client()

    prompt = f"""Write a tailored cover letter for this job posting.

JOB POSTING:
Company: {company}
Title: {title}

{job_description}

Output exactly 11 lines as specified in your instructions. Nothing more."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        system=COVER_SYSTEM,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


# ---------------------------------------------------------------------------
# Document writing — format-preserving
# ---------------------------------------------------------------------------

def _copy_run_format(source_run, target_run) -> None:
    """Copy font formatting from one run to another."""
    target_run.bold          = source_run.bold
    target_run.italic        = source_run.italic
    target_run.underline     = source_run.underline
    target_run.font.size     = source_run.font.size
    target_run.font.name     = source_run.font.name
    target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color and source_run.font.color.type else None


def write_docx_from_template(text: str, output_path: Path, template_path: Path) -> None:
    """
    Write tailored text to a .docx while preserving the template's formatting.

    Strategy:
      - Copy the template document as the base
      - Walk through paragraphs in order, replacing text content
        while keeping paragraph styles, fonts, and spacing intact
      - Extra paragraphs from Claude are appended using the style
        of the last matched paragraph
    """
    import shutil
    from docx import Document

    # Start from a copy of the master so all styles are inherited
    shutil.copy2(str(template_path), str(output_path))
    doc = Document(str(output_path))

    # Split Claude's output into non-empty lines
    new_lines = [l for l in text.split('\n') if l.strip()]
    template_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    for i, para in enumerate(template_paragraphs):
        if i >= len(new_lines):
            # Clear any remaining template paragraphs we didn't fill
            for run in para.runs:
                run.text = ''
            continue

        new_text = new_lines[i]

        # Preserve the first run's formatting, clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''
        else:
            para.clear()
            para.add_run(new_text)

    # If Claude returned more lines than the template has paragraphs,
    # append extras using the style of the last template paragraph
    if len(new_lines) > len(template_paragraphs):
        last_style = template_paragraphs[-1].style if template_paragraphs else 'Normal'
        for extra_line in new_lines[len(template_paragraphs):]:
            p = doc.add_paragraph(extra_line, style=last_style)

    # Remove any leftover bullet/customization guide paragraphs
    for para in doc.paragraphs:
        if para.style.name == 'List Bullet' and not para.text.strip():
            p = para._element
            p.getparent().remove(p)

    doc.save(str(output_path))
    

# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def tailor_for_job(
    company: str,
    title: str,
    url: str = "",
    pasted_text: str = ""
) -> dict:
    """
    Full tailoring pipeline for a single job posting.

    Args:
        company:     Company name (from DB or user input)
        title:       Job title (from DB or user input)
        url:         Job posting URL (tried first)
        pasted_text: Fallback if URL fetch fails

    Returns:
        {
            "success":       bool,
            "resume_path":   str,
            "cover_path":    str,
            "output_dir":    str,
            "fetch_status":  str,   # 'url', 'pasted', 'failed'
            "error":         str,   # empty if success
        }
    """
    result = {
        "success":      False,
        "resume_path":  "",
        "cover_path":   "",
        "output_dir":   "",
        "fetch_status": "",
        "error":        "",
    }

    # --- Validate master files exist ---
    if not MASTER_RESUME.exists():
        result["error"] = f"Master resume not found at {MASTER_RESUME}"
        return result
    if not MASTER_COVER.exists():
        result["error"] = f"Master cover letter not found at {MASTER_COVER}"
        return result

    # --- Get job description ---
    job_description = ""

    if url:
        fetched, status = fetch_job_description(url)
        if status == "ok" and fetched:
            job_description = fetched
            result["fetch_status"] = "url"
        else:
            result["fetch_status"] = f"url_failed({status})"

    if not job_description and pasted_text and pasted_text.strip():
        job_description = pasted_text.strip()
        result["fetch_status"] = "pasted"

    if not job_description:
        result["error"] = (
            "Could not retrieve job description. "
            "URL fetch failed and no pasted text was provided."
        )
        result["fetch_status"] = "failed"
        return result

    # --- Read master documents ---
    try:
        master_resume_text = read_docx_text(MASTER_RESUME)
        master_cover_text  = read_docx_text(MASTER_COVER)
    except Exception as e:
        result["error"] = f"Failed to read master documents: {e}"
        return result

    # --- Create output directory ---
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_') or "Unknown_Company"
    today        = date.today().strftime("%Y-%m-%d")
    job_dir      = OUTPUT_DIR / f"{safe_company}_{today}"
    job_dir.mkdir(parents=True, exist_ok=True)

    # --- Tailor resume ---
    try:
        tailored_resume = tailor_resume_with_claude(
            master_resume_text, job_description, company, title
        )
    except Exception as e:
        result["error"] = f"Claude resume tailoring failed: {e}"
        return result

    # --- Tailor cover letter ---
    try:
        tailored_cover = tailor_cover_letter_with_claude(
            master_cover_text, job_description, company, title
        )
    except Exception as e:
        result["error"] = f"Claude cover letter tailoring failed: {e}"
        return result

    # --- Save documents ---
    resume_filename = f"Cindy_Keller_Resume_{safe_company}.docx"
    cover_filename  = f"Cindy_Keller_CoverLetter_{safe_company}.docx"
    resume_path     = job_dir / resume_filename
    cover_path      = job_dir / cover_filename

    try:
        write_docx_from_template(tailored_resume, resume_path, MASTER_RESUME)
        write_docx_from_template(tailored_cover,  cover_path,  MASTER_COVER)
    except Exception as e:
        result["error"] = f"Failed to save documents: {e}"
        return result

    result.update({
        "success":     True,
        "resume_path": str(resume_path),
        "cover_path":  str(cover_path),
        "output_dir":  str(job_dir),
    })
    return result
